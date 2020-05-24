from django.db.models import Q
from django.utils.crypto import get_random_string
from push_notifications.models import GCMDevice
from django.db.models.signals import pre_save , post_save
from docx import Document
from docx.shared import Inches , RGBColor , Pt
from docx.enum.text import WD_LINE_SPACING

def search(queryset , search_dic , serializer , type):
#    first_search , *searchs = search_dic
    filter_field , *filters = search_dic.keys()
    search_field , *fields = search_dic.values()

    if search_field is not None and filter_field is not None:
        q_objects = Q()
        q_objects = Q(**{ filter_field + '__icontains' : search_field[0] })

        if len(search_dic) > 1 :

            if type == "OR":
                for filter , field in zip(filters, fields):
                    q_objects |= Q(**{ filter + '__icontains' : field[0]})
            elif type == "AND":
                for filter , field in zip(filters, fields):
                    q_objects &= Q(**{ filter + '__icontains' : field[0]})

        found_elements = queryset.filter(q_objects)
        resultlist = []

        for found in found_elements:
            serializers = serializer(found)
            resultlist.append(serializers.data)

        return resultlist
    else:
        return  "please provide a valid search url"



def get_or_create_model_instance(arguments_list , arguments_values , model_manager , model , search_field , value):

    model_list = model_manager.objects.filter(**{search_field: value})

    if not model_list:
        i = 0
        for argument_l in arguments_list:
            setattr(model, argument_l , arguments_values[i])
            i += 1
            try:
                model.save()
            except ValueError as err:
                return {"message" : err.args}
    else:
        model = model_list[0]

    return model


def get_temporal_password(user):
    unique_id = get_random_string(length=32)
    user.temporal_password = unique_id
    user.save()
    return unique_id


def send_notification(user , title ,  message):
    devices = GCMDevice.objects.all().filter(user = user)
    for device in devices:
        device.send_message(message, extra={"title": title , "icon": "ic_leaf"})
    #device = GCMDevice.objects.get(registration_id=gcm_reg_id)


def save_image_url(sender, instance, **kwargs):
    post_save.disconnect(save_image_url, sender=sender)
    instance.photo_url = instance.photo.url
    post_save.connect(save_image_url, sender=sender)


#its relatively dirty
def build_report(specimen):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(16)
    
    heading_style = document.styles['Heading 1']
    heading_style.font.size = Pt(18)

    heading = document.add_heading()
    heading.alignment = 1

    run = heading.add_run("Herbario Nacional de Nicaragua", 0)
    font = run.font
    font.color.rgb = RGBColor(0, 0, 0)
 
    paragraph = document.add_paragraph(specimen.species.genus.name.upper())
    paragraph.paragraph_format.space_before = Pt(40)

    document.add_paragraph(specimen.species.scientific_name)
    document.add_paragraph("Pais: {}".format(specimen.city.state.country.name))

    document.add_paragraph("Departamento de: {}, Municipio {} ".format(specimen.city.state.name , specimen.city.name , specimen.location))

    document.add_paragraph("{}, {}".format(specimen.species.common_name, specimen.description))

    paragraph = document.add_paragraph("{} {}      {}      {}".format(specimen.user.first_name , specimen.user.last_name, specimen.pk , specimen.date_received.strftime("%d/%m/%Y")))
    paragraph.alignment = 1

    heading = document.add_heading()
    heading.alignment = 1
    heading.paragraph_format.space_before = Pt(40)
    heading.paragraph_format.space_after = 0
    run = heading.add_run("Facultad de Ciencias, Tecnologia y Ambiente", 0)
    font = run.font
    font.color.rgb = RGBColor(0, 0, 0)

    heading = document.add_heading()
    heading.alignment = 1
    run = heading.add_run("Universidad Centroamericana UCA", 0)
    heading.paragraph_format.space_before = 0
    font = run.font
    font.color.rgb = RGBColor(0, 0, 0)

    document.add_page_break()
    return document
    